"""PDF 圖片提取與 OCR 文字辨識模組。"""

import io
from os import mkdir, path
import sys
from datetime import datetime
from glob import glob
import configparser
import shutil
import fitz
from PIL import Image
import easyocr
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENTATION


# determine if application is a script file or frozen exe
if getattr(sys, 'frozen', False):
    application_path = path.dirname(sys.executable)
    config = configparser.ConfigParser()
    config_path = path.join(application_path, '_internal', 'settings.cfg')
    config.read(config_path)
    PROCESSED_FILES_PATH = path.join(
        application_path,
        config['Files']['PROCESSED_FILES_PATH'])
    TEMP_PATH = path.join(
        application_path, '_internal', config['Files']['TEMP_PATH'])
    MODEL_PATH = path.join(
        application_path,
        '_internal',
        config['OCR']['MODEL_PATH'])
    USER_NETWORK_PATH = path.join(
        application_path,
        '_internal',
        config['OCR']['USER_NETWORK_PATH'])
    GPU_ENABLED = bool(config['OCR']['GPU_ENABLED'])
else:  # elif __file__:
    application_path = path.dirname(__file__)
    config = configparser.ConfigParser()
    config_path = path.join(application_path, 'settings.cfg')
    config.read(config_path)
    PROCESSED_FILES_PATH = path.join(
        application_path,
        config['Files']['PROCESSED_FILES_PATH'])
    TEMP_PATH = path.join(application_path, config['Files']['TEMP_PATH'])
    MODEL_PATH = path.join(application_path, config['OCR']['MODEL_PATH'])
    USER_NETWORK_PATH = path.join(
        application_path,
        config['OCR']['USER_NETWORK_PATH'])
    GPU_ENABLED = bool(config['OCR']['GPU_ENABLED'])

reader = easyocr.Reader(
    ['ch_tra', 'en'],
    model_storage_directory=MODEL_PATH,
    user_network_directory=USER_NETWORK_PATH,
    gpu=GPU_ENABLED,
    download_enabled=True)


def save_pdf_images(pdf_file_path: str = None):
    ''' 將 pdf 檔案中的圖片存檔

    Args:
        pdf_file_path (str): pdf 檔案的路徑
    '''
    pdf_file = fitz.open(pdf_file_path)
    for page_number, page in enumerate(pdf_file):

        # 取得圖片並存檔
        for image_index, img in enumerate(page.get_images(), start=1):
            xref = img[0]
            # extract image bytes
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]
            # get image extension
            image_ext = base_image["ext"]

            # Create a PIL Image object from the image bytes
            pil_image = Image.open(io.BytesIO(image_bytes))

            # Save the image to disk
            image_path = path.join(
                TEMP_PATH,
                f"image_{page_number}_{image_index}.{image_ext}")
            pil_image.save(image_path)


def string_expect_case(processed_string: str = None):
    ''' 處理預期錯誤的文字

    Args:
        processed_string (str): OCR 辨識後的文字
    '''
    final = processed_string
    final = final.replace(" ", "")  # 去除空白
    final = final.replace("+", "十")  # 將 + 轉換為 十
    final = final.replace("芸", "巷")  # 將 芸 轉換為 巷
    final = final.replace("臺北市中王區", "臺北市中正區")  # 將 臺北市中王區 轉換為 臺北市中正區
    final = final.replace("重蔓南路", "重慶南路")  # 將 重蔓南路 轉換為 重慶南路
    return final


def ocr_task_images():
    ''' 處理 OCR 任務

    Returns:
        str: 處理後的文字
    '''
    final_string = ""
    for image_path in glob(path.join(TEMP_PATH, '*')):
        result = reader.readtext(image_path)
        for (_, text, _) in result:
            final_string = final_string + string_expect_case(text)
    return final_string


def process_ocr_task():
    ''' 批次處理 OCR 任務

    Returns:
        list: 處理後的文字
    '''
    addr_list = []
    for file_full_path in glob(path.join(application_path, '*.pdf')):
        if not path.isdir(TEMP_PATH):
            mkdir(TEMP_PATH)
        if not path.isdir(PROCESSED_FILES_PATH):
            mkdir(PROCESSED_FILES_PATH)

        save_pdf_images(file_full_path)
        final_string = ocr_task_images()
        print(final_string)
        addr_list.append(final_string)
        # 處理完畢，移動檔案
        shutil.rmtree(TEMP_PATH)
        shutil.move(
            file_full_path,
            path.join(PROCESSED_FILES_PATH, path.basename(file_full_path)))
    return addr_list


def docx_task(addr_list: list = None):
    ''' 處理 DOCX 任務 '''
    if addr_list is None:
        return False
    if len(addr_list) == 0:
        return False

    date_now = datetime.now().strftime('%Y%m%d%H%M%S')
    docx_path = path.join(application_path, f'result_{date_now}.docx')

    document = Document()
    document.add_heading('光學辨識結果', 0)
    for addr in addr_list:
        p = document.add_paragraph().add_run(addr)
        p.font.size = Pt(23)
        p.font.name = '微軟正黑體'
        p.font.color.rgb = RGBColor(0, 0, 0)
        p.font.bold = True

    for section in document.sections:
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
    document.save(docx_path)
    return True


def main():
    ''' 主程式 '''
    addr_list = process_ocr_task()
    process_status = docx_task(addr_list)
    if process_status:
        print(f'處理完成，結果儲存於 {application_path}')
    else:
        print('處理失敗')


if __name__ == "__main__":
    main()
